# tools/dokumen.py
import os
import re
import shutil
from datetime import datetime
from werkzeug.utils import secure_filename
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
import docx

from langchain_core.documents import Document

from db_utils import get_db_connection, user_sessions
from tools.tools_rag import index_file
from tools.tools_researcher import is_existing_site

# === Konstanta folder dan tabel ===
UPLOAD_FOLDER = "uploaded_files"
TABLE_NAME = "catatan_site"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# === Fungsi bantu ===
def is_upload_intent(text: str) -> bool:
    upload_keywords = ["upload", "unggah", "kirim", "masukkan", "simpan"]
    text = text.lower()
    return any(kw in text for kw in upload_keywords)

def extract_text_from_file(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".txt":
        return extract_text_from_txt(filepath)
    elif ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    else:
        return f"[❌ Format file '{ext}' tidak didukung]"
def extract_text_from_txt(filepath):
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        with open(filepath, "r", encoding="latin-1") as f:
            text = f.read()

    # 🔧 Normalize newline dan hapus karakter aneh
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u200b", "").replace("\xa0", " ")
    return text

def extract_text_from_pdf(filepath):
    try:
        doc = fitz.open(filepath)
        text = "\n".join(page.get_text() for page in doc)
        return text.strip()
    except Exception as e:
        return f"[❌ Gagal ekstrak PDF: {e}]"

def extract_text_from_docx(filepath):
    try:
        doc = docx.Document(filepath)
        full_text = []

        # 1. Paragraf biasa
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # 2. Tabel (misalnya kalau user ketik isi notulensi dalam tabel)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))  # Optional: pakai pemisah antar cell

        return "\n".join(full_text)

    except Exception as e:
        return f"[❌ Gagal ekstrak DOCX: {e}]"

def parse_general_blocks(site_name: str, text: str, tanggal: str, jam: str, file_path: str, original_filename: str, custom_name=None, file_type=None):
    conn = get_db_connection()
    cursor = conn.cursor()
    success = 0

    blocks = re.split(r"\n\s*\n", text.strip())  # Pisahkan blok berdasarkan newline ganda

    for block in blocks:
        lines = block.strip().splitlines()
        isi = []
        status = "aktif"
        tanggal_selesai = None

        # Abaikan blok yang merupakan heading umum
        if any(re.search(r"notulensi site", l, re.IGNORECASE) for l in lines):
            continue

        for line in lines:
            line = line.strip()
            if re.match(r"^status\s*:?", line, re.IGNORECASE):
                status = re.sub(r"^status\s*:?", "", line, flags=re.IGNORECASE).strip().lower()
            elif re.match(r"^tanggal selesai\s*:?", line, re.IGNORECASE):
                tanggal_selesai = re.sub(r"^tanggal selesai\s*:?", "", line, flags=re.IGNORECASE).strip()
            else:
                isi.append(line)

        if isi:
            isi_catatan = " ".join(isi)
            cursor.execute(f"""
                INSERT INTO {TABLE_NAME}
                (site_name, tanggal, jam, isi_catatan, file_path, original_filename, custom_name, file_type, status, tanggal_selesai)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                site_name, tanggal, jam, isi_catatan, file_path,
                original_filename, custom_name, file_type, status, tanggal_selesai
            ))
            success += 1

    conn.commit()
    cursor.close()
    conn.close()
    return f"✅ {success} catatan berhasil disimpan dari blok-blok umum."


def parse_and_save_to_db(text, site_name, file_path=None, original_filename=None):
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    parsed_indexes = set()

    catatan_list = []
    curr_tanggal = None
    curr_jam = None
    curr_status = None
    curr_tanggal_selesai = None
    buffer_isi = []

    def normalize_isi(isi_text):
        isi_text = isi_text.strip()
        if isi_text.lower().startswith("isi:"):
            isi_text = isi_text[4:].strip()
        return isi_text

    def simpan_catatan():
        nonlocal catatan_list, curr_tanggal, curr_jam, curr_status, curr_tanggal_selesai, buffer_isi
        if buffer_isi:
            catatan_list.append({
                "tanggal": curr_tanggal,
                "jam": curr_jam,
                "status": curr_status or "aktif",
                "isi_catatan": normalize_isi(" ".join(buffer_isi)),
                "tanggal_selesai": curr_tanggal_selesai
            })
        buffer_isi.clear()
        curr_status = None
        curr_tanggal_selesai = None

    # 🧠 Tangani format emoji (multi-line dan satu baris)
    i = 0
    while i < len(lines):
        if i in parsed_indexes:
            i += 1
            continue

        line = lines[i]
        # Multi-line emoji format
        if "📅" in line and "⏰" in line and i + 2 < len(lines):
            line2 = lines[i + 1]
            line3 = lines[i + 2]
            if "✅" in line2 and "📌" in line3:
                tanggal_match = re.search(r"📅\s*(.*?)\s*⏰\s*([\d:]+)", line)
                isi = line2.replace("✅", "").strip()
                tanggal_selesai_match = re.search(r"📌\s*Tanggal Selesai[:：]?\s*(.*)", line3, flags=re.IGNORECASE)

                if tanggal_match and tanggal_selesai_match:
                    tanggal = tanggal_match.group(1).strip()
                    jam = tanggal_match.group(2).strip()
                    tanggal_selesai = tanggal_selesai_match.group(1).strip()
                    catatan_list.append({
                        "tanggal": tanggal,
                        "jam": jam,
                        "isi_catatan": isi,
                        "status": "selesai",
                        "tanggal_selesai": tanggal_selesai
                    })
                    parsed_indexes.update({i, i + 1, i + 2})
                    i += 3
                    continue
        i += 1

    # Multiline ⏳ format
    i = 0
    while i < len(lines):
        if i in parsed_indexes:
            i += 1
            continue

        line = lines[i]
        if "📅" in line and "⏰" in line:
            tanggal_match = re.search(r"📅\s*(.*?)\s*⏰\s*([\d:]+)", line)
            if tanggal_match:
                tanggal = tanggal_match.group(1).strip()
                jam = tanggal_match.group(2).strip()
                j = i + 1
                while j < len(lines) and lines[j].startswith("⏳"):
                    isi = lines[j].replace("⏳", "").strip()
                    catatan_list.append({
                        "tanggal": tanggal,
                        "jam": jam,
                        "isi_catatan": isi,
                        "status": "aktif",
                        "tanggal_selesai": None
                    })
                    parsed_indexes.add(j)
                    j += 1
                parsed_indexes.add(i)
                i = j
                continue
        i += 1

    # Emoji satu baris
    emoji_line_regex = re.compile(
        r"📅\s*(.+?)\s*⏰\s*([\d:]+)\s*✅\s*(.+?)\s*📌\s*Tanggal Selesai[:：]?\s*(.+)",
        flags=re.IGNORECASE
    )
    for idx, line in enumerate(lines):
        if idx in parsed_indexes:
            continue

        match = emoji_line_regex.search(line.replace("\u200b", "").replace("\xa0", " "))
        if match:
            tanggal = match.group(1).strip()
            jam = match.group(2).strip()
            isi_catatan = match.group(3).strip()
            tanggal_selesai = match.group(4).strip()
            catatan_list.append({
                "tanggal": tanggal,
                "jam": jam,
                "isi_catatan": isi_catatan,
                "status": "selesai",
                "tanggal_selesai": tanggal_selesai
            })
            parsed_indexes.add(idx)

    # Format umum: label biasa
    for idx, line in enumerate(lines):
        if idx in parsed_indexes:
            continue

        line = line.strip()
        line_lc = line.lower()

        if re.match(r"^(📝\s*)?notulensi site", line_lc):
            continue

        if line_lc.startswith("tanggal selesai"):
            curr_tanggal_selesai = line.split(":", 1)[-1].strip()
        elif line_lc.startswith("tanggal") and "jam" in line_lc:
            simpan_catatan()
            tanggal_part = re.search(r"Tanggal:\s*(.*?)\s*(?=Jam:)", line, re.IGNORECASE)
            jam_part = re.search(r"Jam:\s*([\d:]+)", line, re.IGNORECASE)
            curr_tanggal = tanggal_part.group(1).strip() if tanggal_part else None
            curr_jam = jam_part.group(1).strip() if jam_part else None
        elif line_lc.startswith("status"):
            curr_status = "selesai" if "selesai" in line_lc or "done" in line_lc else "aktif"
        elif line_lc.startswith("isi:"):
            buffer_isi.append(line.split(":", 1)[-1].strip())
        else:
            buffer_isi.append(line.strip())

    simpan_catatan()

    if not catatan_list:
        return None

    # ✅ Clean dan validasi sebelum masukkan ke DB
    def clean_date(raw):
        if not raw:
            return None
        return raw.strip().rstrip("-").strip()
    
    success = 0
    with get_db_connection() as conn:
        with conn.cursor() as cur:
            for c in catatan_list:
                tanggal = clean_date(c["tanggal"])
                tanggal_selesai = clean_date(c["tanggal_selesai"])

                cur.execute(f"""
                    INSERT INTO {TABLE_NAME}
                    (site_name, tanggal, jam, isi_catatan, file_path, original_filename, status, tanggal_selesai)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    site_name, tanggal, c["jam"], c["isi_catatan"],
                    file_path, original_filename, c["status"], tanggal_selesai
                ))
                success += 1
            conn.commit()

    return f"✅ {success} catatan berhasil disimpan ke DB untuk site {site_name.upper()}"

def unggah_dokumen(query: str, user_id="default") -> str:
    if not is_upload_intent(query):
        return "⚠️ Jika ingin unggah dokumen, gunakan kata seperti *upload* atau *unggah* di kalimat kamu."

    match = re.search(r"\bsite\s+([\w\-]+)", query, re.IGNORECASE)
    if not match:
        return "⚠️ Harap sebutkan nama site, contoh: Saya ingin upload dokumen site cilacap_pl."

    site = match.group(1).strip().lower()
    if not is_existing_site(site):
        return f"❌ Site {site} tidak dikenali. Pastikan site tersebut valid dan sudah terdaftar."

    user_sessions[user_id] = site
    return f"📤 Silakan unggah dokumen untuk site **{site.upper()}** melalui box upload di bawah ini."

def simpan_file(file, user_id="default", custom_name=None):
    if file is None:
        return "⚠️ Harap pilih file untuk diunggah."

    site_name = user_sessions.get(user_id)
    if not site_name:
        return "⚠️ Harap ketikkan nama site terlebih dahulu sebelum upload."

    site_name = site_name.strip().lower()
    original_filename = secure_filename(os.path.basename(file.name))
    file_ext = os.path.splitext(original_filename)[1].lower()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_filename = f"{site_name}{file_ext}"

    user_folder = os.path.join(UPLOAD_FOLDER, user_id)
    os.makedirs(user_folder, exist_ok=True)
    destination_path = os.path.join(user_folder, safe_filename)

    tanggal = datetime.now().strftime("%A, %d %B %Y")
    jam = datetime.now().strftime("%H:%M:%S")
    nama_display = custom_name.strip() if custom_name and custom_name.strip() else None

    try:
        shutil.copy(file.name, destination_path)

        isi_catatan = None
        if file_ext in [".jpg", ".jpeg", ".png"]:
            image = Image.open(destination_path)
            isi_catatan = pytesseract.image_to_string(image).strip()
            if isi_catatan:
                doc = Document(page_content=isi_catatan, metadata={"source": safe_filename})
                index_file([doc])
        elif file_ext in [".txt", ".pdf", ".docx"]:
            isi_catatan = extract_text_from_file(destination_path)
            index_file(file_path=destination_path, site_name=site_name)

            hasil_parse = parse_and_save_to_db(
                text=isi_catatan,
                site_name=site_name,
                file_path=safe_filename,
                original_filename=original_filename
             )
            if hasil_parse:
                return hasil_parse
            
            # === 1. Deteksi baris-baris dengan format emoji satu baris ===
            lines = isi_catatan.split("\n")
            success_emoji = 0

            for line in lines:
                emoji_line_match = re.search(
                    r"📅\s*(.+?)\s*⏰\s*([\d:]+)\s*✅\s*(.+?)\s*📌\s*Tanggal Selesai[:：]?\s*(.+)",
                    line.strip(), flags=re.IGNORECASE
                )
                if emoji_line_match:
                    tgl = emoji_line_match.group(1).strip()
                    jm = emoji_line_match.group(2).strip()
                    isi = emoji_line_match.group(3).strip()
                    tgl_selesai = emoji_line_match.group(4).strip()

                    with get_db_connection() as conn:
                        with conn.cursor() as cur:
                            cur.execute(f"""
                            INSERT INTO {TABLE_NAME}
                            (site_name, tanggal, jam, isi_catatan, file_path, original_filename, custom_name, file_type, status, tanggal_selesai)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                            """, (
                            site_name, tgl, jm, isi, safe_filename,
                            original_filename, nama_display, file_ext[1:], "selesai", tgl_selesai
                            ))
                        conn.commit()
                        success_emoji += 1

            if success_emoji > 0:
                return f"✅ {success_emoji} catatan berhasil disimpan dari format emoji."
            
            # === 2. Coba parse format structured ===
            catatan_terstruktur = False
            if "Tanggal:" in isi_catatan and "Status:" in isi_catatan:
                hasil_parse = parse_and_save_to_db(
                    text=isi_catatan,
                    site_name=site_name,
                    file_path=safe_filename,
                    original_filename=original_filename
                )
                if hasil_parse:
                    return hasil_parse
                else:
                    catatan_terstruktur = False

            # === 3. Coba parse blok umum ===
            if isi_catatan:
                return parse_general_blocks(
                    site_name=site_name,
                    text=isi_catatan,
                    tanggal=tanggal,
                    jam=jam,
                    file_path=safe_filename,
                    original_filename=original_filename,
                    custom_name=nama_display,
                    file_type=file_ext[1:]
                )

        return f"✅ File {original_filename} berhasil disimpan dan dicatat."

    except Exception as e:
        return f"❌ Gagal menyimpan file: {e}"


